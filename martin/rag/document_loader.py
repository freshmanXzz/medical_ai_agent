"""
DocumentLoader - 文档加载器
支持医学知识库文档的解析和切分
"""
import os
import csv
import re
from typing import List, Dict, Optional
from dataclasses import dataclass

# 导入统一日志工具
from martin.util import AppLogger

logger = AppLogger.setup_logging(__name__)

# 可选依赖：PDF 和 Word 解析
try:
    from pypdf import PdfReader
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    from docx import Document
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

# 常量定义
DEFAULT_CHUNK_SIZE = 500
DEFAULT_CHUNK_OVERLAP = 50


@dataclass
class DocumentChunk:
    """文档切分数据类"""
    content: str
    source: str
    category: str
    chunk_index: int
    metadata: Dict


class DocumentLoader:
    """
    文档加载器
    
    支持Markdown和CSV文档的解析和切分
    
    Args:
        chunk_size: 切分大小（字符数）
        chunk_overlap: 切分重叠（字符数）
    """
    
    def __init__(self, chunk_size: int = DEFAULT_CHUNK_SIZE, chunk_overlap: int = DEFAULT_CHUNK_OVERLAP):
        """
        初始化文档加载器
        
        Args:
            chunk_size: 切分大小（字符数），默认500
            chunk_overlap: 切分重叠（字符数），默认50
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        logger.info(f"DocumentLoader初始化完成，chunk_size={chunk_size}, chunk_overlap={chunk_overlap}")

    def load_pdf(self, filepath: str, category: str = None) -> List[DocumentChunk]:
        """
        加载PDF文档

        Args:
            filepath: 文件路径
            category: 分类，默认为文件名（不含扩展名）

        Returns:
            文档切分列表

        Raises:
            FileNotFoundError: 文件不存在
            ValueError: 文件格式不支持或缺少依赖
        """
        if not os.path.exists(filepath):
            logger.error(f"文件不存在: {filepath}")
            raise FileNotFoundError(f"文件不存在: {filepath}")

        if not filepath.lower().endswith('.pdf'):
            logger.error(f"不支持的文件格式: {filepath}")
            raise ValueError(f"不支持的文件格式，仅支持.pdf文件: {filepath}")

        if not HAS_PYPDF:
            logger.error("缺少pypdf依赖，请安装: pip install pypdf")
            raise ValueError("缺少pypdf依赖，请安装: pip install pypdf")

        logger.info(f"开始加载PDF文件: {filepath}")

        # 读取PDF内容
        reader = PdfReader(filepath)
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        content = "\n".join(text_parts)
        content = self._clean_text(content)

        # 提取元数据
        filename = os.path.basename(filepath)
        if category is None:
            category = os.path.splitext(filename)[0]

        metadata = {
            'source': filepath,
            'filename': filename,
            'file_type': 'pdf',
            'category': category,
            'file_size': os.path.getsize(filepath),
            'pages': len(reader.pages)
        }

        # 切分文本
        chunks = self._split_text(content)

        # 创建DocumentChunk列表
        document_chunks = []
        for idx, chunk_content in enumerate(chunks):
            chunk_metadata = metadata.copy()
            chunk_metadata['chunk_size'] = len(chunk_content)
            document_chunks.append(DocumentChunk(
                content=chunk_content,
                source=filepath,
                category=category,
                chunk_index=idx,
                metadata=chunk_metadata
            ))

        logger.info(f"PDF文件加载完成: {filepath}, 共生成{len(document_chunks)}个切分")
        return document_chunks

    def load_word(self, filepath: str, category: str = None) -> List[DocumentChunk]:
        """
        加载Word文档（.docx）

        Args:
            filepath: 文件路径
            category: 分类，默认为文件名（不含扩展名）

        Returns:
            文档切分列表

        Raises:
            FileNotFoundError: 文件不存在
            ValueError: 文件格式不支持或缺少依赖
        """
        if not os.path.exists(filepath):
            logger.error(f"文件不存在: {filepath}")
            raise FileNotFoundError(f"文件不存在: {filepath}")

        if not filepath.lower().endswith('.docx'):
            logger.error(f"不支持的文件格式: {filepath}")
            raise ValueError(f"不支持的文件格式，仅支持.docx文件: {filepath}")

        if not HAS_PYTHON_DOCX:
            logger.error("缺少python-docx依赖，请安装: pip install python-docx")
            raise ValueError("缺少python-docx依赖，请安装: pip install python-docx")

        logger.info(f"开始加载Word文件: {filepath}")

        # 读取Word内容
        doc = Document(filepath)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        content = "\n".join(text_parts)
        content = self._clean_text(content)

        # 提取元数据
        filename = os.path.basename(filepath)
        if category is None:
            category = os.path.splitext(filename)[0]

        metadata = {
            'source': filepath,
            'filename': filename,
            'file_type': 'word',
            'category': category,
            'file_size': os.path.getsize(filepath),
            'paragraphs': len(doc.paragraphs)
        }

        # 切分文本
        chunks = self._split_text(content)

        # 创建DocumentChunk列表
        document_chunks = []
        for idx, chunk_content in enumerate(chunks):
            chunk_metadata = metadata.copy()
            chunk_metadata['chunk_size'] = len(chunk_content)
            document_chunks.append(DocumentChunk(
                content=chunk_content,
                source=filepath,
                category=category,
                chunk_index=idx,
                metadata=chunk_metadata
            ))

        logger.info(f"Word文件加载完成: {filepath}, 共生成{len(document_chunks)}个切分")
        return document_chunks

    def load_markdown(self, filepath: str, category: str = None) -> List[DocumentChunk]:
        """
        加载Markdown文档
        
        Args:
            filepath: 文件路径
            category: 分类，默认为文件名（不含扩展名）
        
        Returns:
            文档切分列表
        
        Raises:
            FileNotFoundError: 文件不存在
            ValueError: 文件格式不支持
        """
        if not os.path.exists(filepath):
            logger.error(f"文件不存在: {filepath}")
            raise FileNotFoundError(f"文件不存在: {filepath}")
        
        if not filepath.lower().endswith('.md'):
            logger.error(f"不支持的文件格式: {filepath}")
            raise ValueError(f"不支持的文件格式，仅支持.md文件: {filepath}")
        
        logger.info(f"开始加载Markdown文件: {filepath}")
        
        # 读取文件内容
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 清理文本
        content = self._clean_text(content)
        
        # 提取元数据
        filename = os.path.basename(filepath)
        if category is None:
            category = os.path.splitext(filename)[0]
        
        metadata = {
            'source': filepath,
            'filename': filename,
            'file_type': 'markdown',
            'category': category,
            'file_size': os.path.getsize(filepath)
        }
        
        # 切分文本
        chunks = self._split_text(content)
        
        # 创建DocumentChunk列表
        document_chunks = []
        for idx, chunk_content in enumerate(chunks):
            chunk_metadata = metadata.copy()
            chunk_metadata['chunk_size'] = len(chunk_content)
            document_chunks.append(DocumentChunk(
                content=chunk_content,
                source=filepath,
                category=category,
                chunk_index=idx,
                metadata=chunk_metadata
            ))
        
        logger.info(f"Markdown文件加载完成: {filepath}, 共生成{len(document_chunks)}个切分")
        return document_chunks
    
    def load_csv(self, filepath: str, category: str = None) -> List[DocumentChunk]:
        """
        加载CSV文档
        
        Args:
            filepath: 文件路径
            category: 分类，默认为文件名（不含扩展名）
        
        Returns:
            文档切分列表
        
        Raises:
            FileNotFoundError: 文件不存在
            ValueError: 文件格式不支持
        """
        if not os.path.exists(filepath):
            logger.error(f"文件不存在: {filepath}")
            raise FileNotFoundError(f"文件不存在: {filepath}")
        
        if not filepath.lower().endswith('.csv'):
            logger.error(f"不支持的文件格式: {filepath}")
            raise ValueError(f"不支持的文件格式，仅支持.csv文件: {filepath}")
        
        logger.info(f"开始加载CSV文件: {filepath}")
        
        # 提取元数据
        filename = os.path.basename(filepath)
        if category is None:
            category = os.path.splitext(filename)[0]
        
        document_chunks = []
        
        with open(filepath, 'r', encoding='utf-8') as f:
            # 使用csv模块读取
            reader = csv.DictReader(f)
            headers = reader.fieldnames
            
            for row_idx, row in enumerate(reader):
                # 将每行数据转换为文本
                row_content = self._csv_row_to_text(row, headers)
                
                # 清理文本
                row_content = self._clean_text(row_content)
                
                # 如果内容过长，进行切分
                if len(row_content) > self.chunk_size:
                    chunks = self._split_text(row_content)
                    for chunk_idx, chunk_content in enumerate(chunks):
                        chunk_metadata = {
                            'source': filepath,
                            'filename': filename,
                            'file_type': 'csv',
                            'category': category,
                            'row_index': row_idx,
                            'headers': headers,
                            'chunk_size': len(chunk_content)
                        }
                        document_chunks.append(DocumentChunk(
                            content=chunk_content,
                            source=filepath,
                            category=category,
                            chunk_index=len(document_chunks),
                            metadata=chunk_metadata
                        ))
                else:
                    chunk_metadata = {
                        'source': filepath,
                        'filename': filename,
                        'file_type': 'csv',
                        'category': category,
                        'row_index': row_idx,
                        'headers': headers,
                        'chunk_size': len(row_content)
                    }
                    document_chunks.append(DocumentChunk(
                        content=row_content,
                        source=filepath,
                        category=category,
                        chunk_index=len(document_chunks),
                        metadata=chunk_metadata
                    ))
        
        logger.info(f"CSV文件加载完成: {filepath}, 共生成{len(document_chunks)}个切分")
        return document_chunks
    
    def load_directory(self, dirpath: str) -> List[DocumentChunk]:
        """
        加载目录下所有文档
        
        Args:
            dirpath: 目录路径
        
        Returns:
            文档切分列表
        
        Raises:
            NotADirectoryError: 路径不是目录
        """
        if not os.path.isdir(dirpath):
            logger.error(f"路径不是目录: {dirpath}")
            raise NotADirectoryError(f"路径不是目录: {dirpath}")
        
        logger.info(f"开始加载目录: {dirpath}")
        
        document_chunks = []
        
        # 遍历目录
        for filename in os.listdir(dirpath):
            filepath = os.path.join(dirpath, filename)

            if os.path.isfile(filepath):
                try:
                    if filename.lower().endswith('.md'):
                        chunks = self.load_markdown(filepath)
                        document_chunks.extend(chunks)
                    elif filename.lower().endswith('.csv'):
                        chunks = self.load_csv(filepath)
                        document_chunks.extend(chunks)
                    elif filename.lower().endswith('.pdf'):
                        chunks = self.load_pdf(filepath)
                        document_chunks.extend(chunks)
                    elif filename.lower().endswith('.docx'):
                        chunks = self.load_word(filepath)
                        document_chunks.extend(chunks)
                    else:
                        logger.warning(f"跳过不支持的文件类型: {filename}")
                except Exception as e:
                    logger.error(f"加载文件失败: {filename}, 错误: {str(e)}")
                    continue
        
        logger.info(f"目录加载完成: {dirpath}, 共生成{len(document_chunks)}个切分")
        return document_chunks
    
    def _split_text(self, text: str) -> List[str]:
        """
        切分文本
        
        使用滑动窗口方式切分文本，保持上下文连续性
        
        Args:
            text: 原文本
        
        Returns:
            切分后的文本列表
        """
        if not text:
            return []
        
        # 如果文本长度小于chunk_size，直接返回
        if len(text) <= self.chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            # 计算当前chunk的结束位置
            end = start + self.chunk_size
            
            # 如果不是最后一个chunk，尝试在句子边界切分
            if end < len(text):
                # 查找最近的句子边界（句号、问号、感叹号、换行符）
                boundary_chars = ['。', '！', '？', '；', '\n', '.', '!', '?', ';']
                best_boundary = -1
                
                for char in boundary_chars:
                    # 在chunk末尾附近查找边界（向后50字符范围内）
                    boundary = text.rfind(char, start + self.chunk_size - 50, end + 50)
                    if boundary > start and boundary < end + 50:
                        if best_boundary == -1 or boundary > best_boundary:
                            best_boundary = boundary
                
                # 如果找到合适的边界，在边界后切分
                if best_boundary > start:
                    end = best_boundary + 1
            
            # 提取chunk
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # 计算下一个chunk的起始位置（考虑重叠）
            next_start = end - self.chunk_overlap
            
            # 确保start不会倒退
            if next_start <= start:
                start = end
            else:
                start = next_start
        
        return chunks
    
    def _clean_text(self, text: str) -> str:
        """
        清理文本
        
        移除多余的空白字符和特殊字符
        
        Args:
            text: 原文本
        
        Returns:
            清理后的文本
        """
        if not text:
            return ""
        
        # 移除多余的空白字符
        text = re.sub(r'\s+', ' ', text)
        
        # 移除特殊控制字符
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        # 移除首尾空白
        text = text.strip()
        
        return text
    
    def _csv_row_to_text(self, row: Dict, headers: List[str]) -> str:
        """
        将CSV行转换为文本
        
        Args:
            row: CSV行数据（字典格式）
            headers: 列标题列表
        
        Returns:
            格式化的文本
        """
        if not row or not headers:
            return ""
        
        parts = []
        for header in headers:
            value = row.get(header, '')
            if value:
                parts.append(f"{header}: {value}")
        
        return " | ".join(parts)